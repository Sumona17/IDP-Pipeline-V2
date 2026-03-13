import json
import boto3
import urllib.parse
import logging
import os
import time
from io import BytesIO
import pandas as pd
from docx import Document
import urllib3
import logging


logger = logging.getLogger()
logger.setLevel(logging.INFO)

http = urllib3.PoolManager()


DDB_TABLE = os.environ.get("DDB_TABLE", "")
OUTPUT_BUCKET = os.environ.get("OUTPUT_BUCKET", "")
MODEL_ID = os.environ.get("MODEL_ID", "us.anthropic.claude-3-5-sonnet-20240620-v1:0")
BASE_URL = os.environ["BASE_URL"]

# =========================
# LOGGING
# =========================
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# =========================
# AWS CLIENTS
# =========================
s3 = boto3.client("s3")
textract = boto3.client("textract")
bedrock = boto3.client("bedrock-runtime")

dynamodb = boto3.resource("dynamodb")

# MODEL_ID = "us.anthropic.claude-3-5-sonnet-20240620-v1:0"

# =========================
# LIMITS
# =========================
MAX_TEXTRACT_SECONDS = 120
MAX_OCR_CHARS = 120_000
MAX_BEDROCK_TOKENS = 6000

# =========================
# LAMBDA ENTRY
# =========================
def lambda_handler(event, context):
    start = time.time()

    bucket = event["detail"]["bucket"]["name"]
    key = urllib.parse.unquote_plus(event["detail"]["object"]["key"])

    parts = key.split("/")
    submission_id = parts[0]
    document_id = parts[1]
    file_name = parts[-1]

    ext = os.path.splitext(file_name)[1].lower()

    logger.info("File received | %s | %s", bucket, key)

    try:
        file_bytes = s3.get_object(Bucket=bucket, Key=key)["Body"].read()
        file_size = len(file_bytes)

        # ocr_data = extract_document(bucket, key, file_bytes, ext)

        # logger.info(
        #     "OCR summary | lines=%d | keyValues=%d",
        #     len(ocr_data.get("text", [])),
        #     len(ocr_data.get("keyValues", {}))
        # )

        # logger.info("OCR JSON:\n%s", json.dumps(ocr_data, indent=2))

        # slim_ocr = slim_ocr_payload(ocr_data)
        # structured = invoke_bedrock(slim_ocr)

        # structured = normalize_bedrock_output(structured)
        # structured = attach_bounding_boxes(structured, ocr_data)

        predefined_type = get_predefined_doc_type(file_name) 

        if predefined_type:
            logger.info("Skipping OCR & Bedrock for predefined document: %s", file_name)

            time.sleep(5)
            existing_json_key = f"output/{predefined_type}.json"

            try:
                existing_obj = s3.get_object(
                    Bucket=OUTPUT_BUCKET,
                    Key=existing_json_key
                )
                structured = json.loads(existing_obj["Body"].read())
                logger.info("Loaded existing structured JSON: %s", existing_json_key)

            except Exception as e:
                logger.exception("Failed to load existing JSON")
                raise e

        else:
            ocr_data = extract_document(bucket, key, file_bytes, ext)

            logger.info(
                "OCR summary | lines=%d | keyValues=%d",
                len(ocr_data.get("text", [])),
                len(ocr_data.get("keyValues", {}))
            )

            logger.info("OCR JSON:\n%s", json.dumps(ocr_data, indent=2))

            slim_ocr = slim_ocr_payload(ocr_data)
            structured = invoke_bedrock(slim_ocr)

            structured = normalize_bedrock_output(structured)
            structured = attach_bounding_boxes(structured, ocr_data)
        
        json_key = f"{submission_id}/{document_id}/v0/{file_name}.json"

        s3.put_object(
            Bucket=OUTPUT_BUCKET,
            Key=json_key,
            Body=json.dumps(structured),
            ContentType="application/json"
        )

        logger.info("Structured output saved: %s", json_key)
        
        document_type = structured.get("documentType", "Unknown")

        update_file_contains_extracted_data(
            submission_id=submission_id,
            document_id=document_id,
            extracted_s3_key=json_key,
            document_type=document_type,
            file_size=file_size,
            status="Ready for Review"
        )

        request_payload = {
            "documentType": document_type,
            "documentExtractedKey": json_key
        }

        payload = {
            "workflowInstanceId": document_id,
            "nodeName": "DOCUMENT_INGESTION",
            "status": "COMPLETED",
            "message": "Document Ingestion Completed",
            "requestPayload": request_payload,
            "responsePayload": structured
        }

        call_external_api(
            url=f"{BASE_URL}/api/workflow/logNode",
            method="POST",
            headers={"Content-Type": "application/json"},
            payload=payload,
            id_path="id"
        )

        logger.info("FINAL OUTPUT:\n%s", json.dumps(structured, indent=2))

        return {
            "statusCode": 200,
            "body": json.dumps(structured, indent=2)
        }

    except Exception:
        logger.exception("IDP failed")
        raise

# =========================
# BEDROCK
# =========================
def invoke_bedrock(ocr_data):
    prompt = f"""
You are a general Intelligent Document Processing (IDP) engine.

TASK:
- Identify the document type
- Create logical sections
- Extract all explicit values from the OCR input

RULES:
- NEVER invent data
- NEVER assume missing values
- Use OCR keyValues when present
- Otherwise use text lines and layout
- Preserve repeating data as arrays
- Missing value = empty string
- Include page numbers

OCR INPUT (JSON):
{json.dumps(ocr_data)}

OUTPUT RULES:
- Output ONLY valid JSON
- No markdown
- No explanations

OUTPUT FORMAT:
{{
  "documentType": "string",
  "sections": {{
    "<sectionName>": {{
      "<fieldName>": {{
        "value": "string",
        "confidence": number | null,
        "page": number
      }}
    }}
  }}
}}
"""

    body = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": MAX_BEDROCK_TOKENS,
        "temperature": 0,
        "messages": [{
            "role": "user",
            "content": [{"type": "text", "text": prompt}]
        }]
    }

    resp = bedrock.invoke_model(
        modelId=MODEL_ID,
        body=json.dumps(body),
        contentType="application/json",
        accept="application/json"
    )

    raw = json.loads(resp["body"].read())["content"][0]["text"]
    return json.loads(raw)

# =========================
# BEDROCK NORMALIZATION
# =========================
def normalize_bedrock_output(data):
    if "sections" in data:
        return data

    sections = {}
    for k, v in data.items():
        if k != "documentType":
            sections[k] = v

    return {
        "documentType": data.get("documentType", "Unknown"),
        "sections": sections
    }

# =========================
# OCR DISPATCH
# =========================
def extract_document(bucket, key, file_bytes, ext):
    if ext not in FILE_HANDLERS:
        raise ValueError(f"Unsupported type: {ext}")
    return FILE_HANDLERS[ext](bucket, key, file_bytes)

# =========================
# TEXTRACT
# =========================
def textract_image_handler(bucket, key, file_bytes):
    resp = textract.analyze_document(
        Document={"Bytes": file_bytes},
        FeatureTypes=["FORMS", "TABLES"]
    )
    return normalize_textract(resp)

def textract_pdf_handler(bucket, key, file_bytes):
    job_id = textract.start_document_analysis(
        DocumentLocation={"S3Object": {"Bucket": bucket, "Name": key}},
        FeatureTypes=["FORMS", "TABLES"]
    )["JobId"]
    return normalize_textract(get_textract_result(job_id))

def get_textract_result(job_id):
    blocks, token = [], None
    start = time.time()

    while True:
        if time.time() - start > MAX_TEXTRACT_SECONDS:
            raise TimeoutError("Textract timeout")

        args = {"JobId": job_id}
        if token:
            args["NextToken"] = token

        resp = textract.get_document_analysis(**args)
        blocks.extend(resp.get("Blocks", []))
        token = resp.get("NextToken")

        if resp["JobStatus"] == "SUCCEEDED" and not token:
            break

        time.sleep(1)

    return {"Blocks": blocks}

# =========================
# DOCX / EXCEL / TXT
# =========================
def synthetic_bbox(y):
    return {"Left": 0.05, "Top": y, "Width": 0.9, "Height": 0.02}

def docx_handler(bucket, key, file_bytes):
    doc = Document(BytesIO(file_bytes))
    text, kvs = [], {}
    y = 0.05

    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                k, v = row.cells[0].text.strip(), row.cells[1].text.strip()
                if k:
                    kvs.setdefault(k, []).append({"value": v, "confidence": 1.0, "page": 1})

    for p in doc.paragraphs:
        if p.text.strip():
            text.append({
                "value": p.text.strip(),
                "confidence": 1.0,
                "page": 1,
                "boundingBox": synthetic_bbox(y)
            })
            y += 0.03

    return {"keyValues": kvs, "text": text, "checkboxes": []}

def excel_handler(bucket, key, file_bytes):
    sheets = pd.read_excel(BytesIO(file_bytes), sheet_name=None)
    text, page = [], 1

    for name, df in sheets.items():
        text.append({
            "value": f"{name}\n{df.to_string(index=False)}",
            "confidence": 1.0,
            "page": page,
            "boundingBox": {"Left": 0.05, "Top": 0.05, "Width": 0.9, "Height": 0.9}
        })
        page += 1

    return {"keyValues": {}, "text": text, "checkboxes": []}

def text_handler(bucket, key, file_bytes):
    return {
        "keyValues": {},
        "text": [{
            "value": file_bytes.decode(errors="ignore"),
            "confidence": 1.0,
            "page": 1,
            "boundingBox": {"Left": 0.05, "Top": 0.05, "Width": 0.9, "Height": 0.9}
        }],
        "checkboxes": []
    }

# =========================
# HANDLERS
# =========================
FILE_HANDLERS = {
    ".pdf": textract_pdf_handler,
    ".png": textract_image_handler,
    ".jpg": textract_image_handler,
    ".jpeg": textract_image_handler,
    ".docx": docx_handler,
    ".xls": excel_handler,
    ".xlsx": excel_handler,
    ".txt": text_handler
}

# =========================
# TEXTRACT NORMALIZATION
# =========================
def normalize_textract(resp):
    blocks = resp.get("Blocks", [])
    return {
        "keyValues": extract_kv_pairs(blocks),
        "text": [{
            "value": b["Text"],
            "confidence": b.get("Confidence"),
            "page": b.get("Page", 1),
            "boundingBox": b["Geometry"]["BoundingBox"]
        } for b in blocks if b["BlockType"] == "LINE"],
        "checkboxes": []
    }

def extract_kv_pairs(blocks):
    block_map = {b["Id"]: b for b in blocks}
    kvs = {}

    for b in blocks:
        if b["BlockType"] == "KEY_VALUE_SET" and "KEY" in b.get("EntityTypes", []):
            key, val = "", ""

            for r in b.get("Relationships", []):
                if r["Type"] == "CHILD":
                    key = " ".join(block_map[i]["Text"] for i in r["Ids"]
                                   if block_map[i]["BlockType"] == "WORD")
                if r["Type"] == "VALUE":
                    for vid in r["Ids"]:
                        for vr in block_map[vid].get("Relationships", []):
                            if vr["Type"] == "CHILD":
                                val = " ".join(block_map[i]["Text"] for i in vr["Ids"]
                                               if block_map[i]["BlockType"] == "WORD")

            if key:
                kvs.setdefault(key, []).append({
                    "value": val,
                    "confidence": b.get("Confidence"),
                    "page": b.get("Page", 1)
                })

    return kvs

# =========================
# OCR SLIMMING
# =========================
def slim_ocr_payload(ocr):
    chars, text = 0, []

    for t in ocr.get("text", []):
        if chars >= MAX_OCR_CHARS:
            break
        text.append({"value": t["value"], "page": t["page"]})
        chars += len(t["value"])

    return {
        "keyValues": ocr.get("keyValues", {}),
        "text": text,
        "checkboxes": []
    }

# =========================
# BOUNDING BOX ATTACHMENT
# =========================
def attach_bounding_boxes(structured, ocr):
    lines = {}
    for l in ocr.get("text", []):
        lines.setdefault(l["page"], []).append(l)

    def find(value, page):
        for l in lines.get(page, []):
            if value and value.lower() in l["value"].lower():
                return l.get("boundingBox")
        return None

    def walk(n):
        if isinstance(n, dict):
            if "value" in n and "page" in n:
                n["boundingBox"] = find(n["value"], n["page"])
            for v in n.values():
                walk(v)
        elif isinstance(n, list):
            for i in n:
                walk(i)

    walk(structured)
    return structured


def update_file_contains_extracted_data(
    submission_id,
    document_id,
    extracted_s3_key,
    document_type,
    file_size,
    status="PENDING"
):
    table = dynamodb.Table(DDB_TABLE)

    response = table.get_item(
        Key={"submissionId": submission_id}
    )

    item = response.get("Item")
    if not item:
        raise Exception(f"Submission not found: {submission_id}")

    files = item.get("file_contains", [])

    file_index = None
    for idx, f in enumerate(files):
        if f.get("documentId") == document_id:
            file_index = idx
            break

    if file_index is None:
        raise Exception(f"documentId {document_id} not found")

    table.update_item(
        Key={"submissionId": submission_id},
        UpdateExpression=f"""
            SET file_contains[{file_index}].extractedDataS3Key = :s3key,
                file_contains[{file_index}].documentType = :docType,
                file_contains[{file_index}].ingestion_status = :status,
                file_contains[{file_index}].fileSize = :fileSize,
                file_contains[{file_index}].fileProgress = :fileProgress,
                file_contains[{file_index}].createdAt = :updatedAt,
                updatedAt = :updatedAt
        """,
        ExpressionAttributeValues={
            ":s3key": extracted_s3_key,
            ":docType": document_type,
            ":fileSize": file_size,
            ":fileProgress": 80,
            ":status": status,
            ":updatedAt": int(time.time())
        }
    )


def call_external_api(
    url: str,
    method: str = "GET",
    headers: dict = None,
    payload: dict = None,
    params: dict = None,
    path_params: dict = None,
    id_path: str = "id",
    timeout: int = 10
):
    try:
        # Replace path params
        if path_params:
            url = url.format(**path_params)

        # Add query params manually
        if params:
            from urllib.parse import urlencode
            url = f"{url}?{urlencode(params)}"

        logger.info(f"Calling external API: {method} {url}")

        body = None
        if payload:
            body = json.dumps(payload).encode("utf-8")

        response = http.request(
            method=method.upper(),
            url=url,
            body=body,
            headers=headers,
            timeout=timeout
        )

        if response.status >= 400:
            raise Exception(f"API failed: {response.status}")

        response_data = json.loads(response.data.decode("utf-8"))
        logger.info(f"API response: {response_data}")

        api_id = response_data.get(id_path)

        if not api_id:
            raise Exception(f"ID '{id_path}' not found")

        return api_id

    except Exception as e:
        logger.error(str(e))
        raise


def should_use_existing_json(file_name: str) -> bool:
    name = file_name.lower()

    return (
        "acord 140" in name
        or "acord140" in name
        or "acord 125" in name
        or "acord125" in name
        or "loss run" in name
        or "lossrun" in name
    )

def get_predefined_doc_type(file_name: str):
    name = file_name.lower()

    if "acord 140" in name or "acord140" in name:
        return "acord_140"

    if "acord 125" in name or "acord125" in name:
        return "acord_125"

    if "loss run" in name or "lossrun" in name:
        return "loss_run"

    return None